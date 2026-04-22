import re
import copy
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FONT_SIZE_MAP = {
    "小初": Pt(36), "一号": Pt(26), "小一": Pt(24),
    "二号": Pt(22), "小二": Pt(18), "三号": Pt(16),
    "小三": Pt(15), "四号": Pt(14), "小四": Pt(12),
    "五号": Pt(10.5), "小五": Pt(9),
}

TEMPLATES = {
    "通用论文": {
        "name": "通用论文模板",
        "description": "适用于大多数高校毕业论文的排版规范",
        "margins": {"top": Cm(2.54), "bottom": Cm(2.54), "left": Cm(3.17), "right": Cm(3.17)},
        "title": {
            "font_cn": "黑体", "font_en": "Times New Roman",
            "size": FONT_SIZE_MAP["小二"], "bold": True,
            "alignment": WD_ALIGN_PARAGRAPH.CENTER,
            "space_before": Pt(0), "space_after": Pt(24),
        },
        "heading1": {
            "font_cn": "黑体", "font_en": "Times New Roman",
            "size": FONT_SIZE_MAP["小三"], "bold": True,
            "alignment": WD_ALIGN_PARAGRAPH.CENTER,
            "space_before": Pt(24), "space_after": Pt(18),
        },
        "heading2": {
            "font_cn": "黑体", "font_en": "Times New Roman",
            "size": FONT_SIZE_MAP["四号"], "bold": True,
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(18), "space_after": Pt(12),
        },
        "heading3": {
            "font_cn": "黑体", "font_en": "Times New Roman",
            "size": FONT_SIZE_MAP["小四"], "bold": True,
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(12), "space_after": Pt(6),
        },
        "body": {
            "font_cn": "宋体", "font_en": "Times New Roman",
            "size": FONT_SIZE_MAP["小四"], "bold": False,
            "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "line_spacing": 1.5,
            "first_line_indent": 2,
            "space_before": Pt(0), "space_after": Pt(0),
        },
        "table": {
            "font_cn": "宋体", "font_en": "Times New Roman",
            "size": FONT_SIZE_MAP["五号"],
            "three_line": True,
        },
        "caption": {
            "font_cn": "宋体", "font_en": "Times New Roman",
            "size": FONT_SIZE_MAP["五号"], "bold": False,
            "alignment": WD_ALIGN_PARAGRAPH.CENTER,
            "line_spacing": 1.5, "first_line_indent": 0,
            "space_before": Pt(6), "space_after": Pt(6),
        },
        "reference": {
            "font_cn": "宋体", "font_en": "Times New Roman",
            "size": FONT_SIZE_MAP["五号"], "bold": False,
            "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "line_spacing": 1.25, "first_line_indent": 0, "hanging_indent": 2,
            "space_before": Pt(0), "space_after": Pt(0),
        },
        "page_number": True,
        "header_footer_size": FONT_SIZE_MAP["小五"],
    },
    "国标期刊": {
        "name": "国标期刊模板",
        "description": "适用于中文核心期刊投稿的排版格式",
        "margins": {"top": Cm(2.5), "bottom": Cm(2.5), "left": Cm(2.5), "right": Cm(2.5)},
        "title": {
            "font_cn": "黑体", "font_en": "Times New Roman",
            "size": FONT_SIZE_MAP["二号"], "bold": True,
            "alignment": WD_ALIGN_PARAGRAPH.CENTER,
            "space_before": Pt(0), "space_after": Pt(30),
        },
        "heading1": {
            "font_cn": "黑体", "font_en": "Times New Roman",
            "size": FONT_SIZE_MAP["四号"], "bold": True,
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(18), "space_after": Pt(12),
        },
        "heading2": {
            "font_cn": "楷体", "font_en": "Times New Roman",
            "size": FONT_SIZE_MAP["小四"], "bold": True,
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(12), "space_after": Pt(6),
        },
        "heading3": {
            "font_cn": "楷体", "font_en": "Times New Roman",
            "size": FONT_SIZE_MAP["小四"], "bold": False,
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(6), "space_after": Pt(6),
        },
        "body": {
            "font_cn": "宋体", "font_en": "Times New Roman",
            "size": FONT_SIZE_MAP["五号"], "bold": False,
            "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "line_spacing": 1.25,
            "first_line_indent": 2,
            "space_before": Pt(0), "space_after": Pt(0),
        },
        "table": {
            "font_cn": "宋体", "font_en": "Times New Roman",
            "size": FONT_SIZE_MAP["小五"],
            "three_line": True,
        },
        "caption": {
            "font_cn": "宋体", "font_en": "Times New Roman",
            "size": FONT_SIZE_MAP["五号"], "bold": False,
            "alignment": WD_ALIGN_PARAGRAPH.CENTER,
            "line_spacing": 1.25, "first_line_indent": 0,
            "space_before": Pt(6), "space_after": Pt(6),
        },
        "reference": {
            "font_cn": "宋体", "font_en": "Times New Roman",
            "size": FONT_SIZE_MAP["五号"], "bold": False,
            "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "line_spacing": 1.25, "first_line_indent": 0, "hanging_indent": 2,
            "space_before": Pt(0), "space_after": Pt(0),
        },
        "page_number": True,
        "header_footer_size": FONT_SIZE_MAP["小五"],
    },
    "简洁商务": {
        "name": "简洁商务模板",
        "description": "适用于公司报告、商务文档的专业排版",
        "margins": {"top": Cm(2.54), "bottom": Cm(2.54), "left": Cm(3.17), "right": Cm(3.17)},
        "title": {
            "font_cn": "微软雅黑", "font_en": "Arial",
            "size": Pt(22), "bold": True,
            "alignment": WD_ALIGN_PARAGRAPH.CENTER,
            "space_before": Pt(0), "space_after": Pt(24),
        },
        "heading1": {
            "font_cn": "微软雅黑", "font_en": "Arial",
            "size": Pt(16), "bold": True,
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(24), "space_after": Pt(12),
        },
        "heading2": {
            "font_cn": "微软雅黑", "font_en": "Arial",
            "size": Pt(14), "bold": True,
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(12), "space_after": Pt(6),
        },
        "heading3": {
            "font_cn": "微软雅黑", "font_en": "Arial",
            "size": Pt(12), "bold": True,
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "space_before": Pt(6), "space_after": Pt(6),
        },
        "body": {
            "font_cn": "微软雅黑", "font_en": "Arial",
            "size": Pt(11), "bold": False,
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "line_spacing": 1.5,
            "first_line_indent": 0,
            "space_before": Pt(0), "space_after": Pt(6),
        },
        "table": {
            "font_cn": "微软雅黑", "font_en": "Arial",
            "size": Pt(9),
            "three_line": False,
        },
        "caption": {
            "font_cn": "微软雅黑", "font_en": "Arial",
            "size": Pt(9), "bold": False,
            "alignment": WD_ALIGN_PARAGRAPH.CENTER,
            "line_spacing": 1.5, "first_line_indent": 0,
            "space_before": Pt(6), "space_after": Pt(6),
        },
        "reference": {
            "font_cn": "微软雅黑", "font_en": "Arial",
            "size": Pt(9), "bold": False,
            "alignment": WD_ALIGN_PARAGRAPH.LEFT,
            "line_spacing": 1.5, "first_line_indent": 0, "hanging_indent": 2,
            "space_before": Pt(0), "space_after": Pt(3),
        },
        "page_number": True,
        "header_footer_size": Pt(9),
    },
}

H1_PATTERNS = [
    r"^第[一二三四五六七八九十百]+章[\s\u3000]",
    r"^摘[\s\u3000]*要$",
    r"^ABSTRACT$",
    r"^参[\s\u3000]*考[\s\u3000]*文[\s\u3000]*献$",
    r"^致[\s\u3000]*谢",
    r"^附[\s\u3000]*录",
    r"^引[\s\u3000]*言$",
    r"^绪[\s\u3000]*论$",
    r"^结[\s\u3000]*论$",
    r"^前[\s\u3000]*言$",
    r"^\d+[\s\u3000]+\S",
]

CONTENT_START_MARKERS = [
    r"^第[一二三四五六七八九十百]+章[\s\u3000]",
    r"^摘[\s\u3000]*要$",
    r"^ABSTRACT$",
    r"^引[\s\u3000]*言$",
    r"^绪[\s\u3000]*论$",
    r"^前[\s\u3000]*言$",
    r"^\d+[\s\u3000]+\S",
]

H2_PATTERNS = [
    r"^\d+\.\d+[\s\u3000]+\S",
    r"^第[一二三四五六七八九十百]+节[\s\u3000]",
]

H3_PATTERNS = [
    r"^\d+\.\d+\.\d+[\s\u3000]+\S",
]

CAPTION_PATTERNS = [
    r"^图[\s\u3000]*\d",
    r"^表[\s\u3000]*\d",
    r"^Figure[\s\u3000]*\d",
    r"^Table[\s\u3000]*\d",
]

REFERENCE_HEADING_RE = re.compile(r"^参[\s\u3000]*考[\s\u3000]*文[\s\u3000]*献$")

_DEFAULT_CAPTION = {
    "font_cn": "宋体", "font_en": "Times New Roman",
    "size": FONT_SIZE_MAP["五号"], "bold": False,
    "alignment": WD_ALIGN_PARAGRAPH.CENTER,
    "line_spacing": 1.5,
    "first_line_indent": 0,
    "space_before": Pt(6), "space_after": Pt(6),
}

_DEFAULT_REFERENCE = {
    "font_cn": "宋体", "font_en": "Times New Roman",
    "size": FONT_SIZE_MAP["五号"], "bold": False,
    "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "line_spacing": 1.25,
    "first_line_indent": 0,
    "hanging_indent": 2,
    "space_before": Pt(0), "space_after": Pt(0),
}


def _is_toc_entry(text):
    """Check if a line looks like a table-of-contents entry (has dot leaders or trailing page numbers)."""
    return bool(
        re.search(r"\.{3,}", text)
        or re.search(r"\u2026{1,}", text)
        or re.search(r"\t+\d+\s*$", text)
        or re.search(r"\s{4,}\d+\s*$", text)
    )


def find_content_start(paragraphs):
    """Locate the first paragraph of actual content, skipping cover / declaration / TOC pages."""
    toc_idx = -1
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        style_name = (para.style.name or "").lower() if para.style else ""
        if re.match(r"^目[\s\u3000]*录$", text) or "toc heading" in style_name:
            toc_idx = i

    if toc_idx >= 0:
        for i in range(toc_idx + 1, len(paragraphs)):
            text = paragraphs[i].text.strip()
            style_name = (paragraphs[i].style.name or "").lower() if paragraphs[i].style else ""
            if not text:
                continue
            if "toc" in style_name:
                continue
            if _is_toc_entry(text):
                continue
            is_content_heading = any(
                re.match(pat, text, re.IGNORECASE) for pat in CONTENT_START_MARKERS
            )
            if is_content_heading:
                return i

    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if not text:
            continue
        for pat in CONTENT_START_MARKERS:
            if re.match(pat, text, re.IGNORECASE):
                return i

    return 0


def detect_role(para, rel_index, total):
    """Detect the structural role of a paragraph. rel_index is relative to content start."""
    text = para.text.strip()
    if not text:
        return "empty"

    style_name = (para.style.name or "").lower() if para.style else ""
    if "title" in style_name:
        return "title"
    if "heading 1" in style_name or style_name == "heading1":
        return "heading1"
    if "heading 2" in style_name or style_name == "heading2":
        return "heading2"
    if "heading 3" in style_name or style_name == "heading3":
        return "heading3"

    if rel_index == 0 and len(text) < 60:
        is_heading = False
        for pat in H1_PATTERNS + H2_PATTERNS + H3_PATTERNS:
            if re.match(pat, text, re.IGNORECASE):
                is_heading = True
                break
        if not is_heading:
            return "title"

    for pat in H3_PATTERNS:
        if re.match(pat, text):
            return "heading3"
    for pat in H2_PATTERNS:
        if re.match(pat, text):
            return "heading2"
    for pat in H1_PATTERNS:
        if re.match(pat, text, re.IGNORECASE):
            return "heading1"

    for pat in CAPTION_PATTERNS:
        if re.match(pat, text, re.IGNORECASE):
            return "caption"

    return "body"


def set_font(run, font_cn, font_en, size, bold=False):
    """Apply font settings to a run, handling both CJK and Latin fonts."""
    run.font.size = size
    run.font.bold = bold
    run.font.name = font_en
    rPr = run.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_cn)


def apply_paragraph_format(para, style_cfg):
    """Apply full paragraph formatting from a style config dict."""
    pf = para.paragraph_format
    pf.alignment = style_cfg["alignment"]
    pf.space_before = style_cfg.get("space_before", Pt(0))
    pf.space_after = style_cfg.get("space_after", Pt(0))

    ls = style_cfg.get("line_spacing")
    if ls is not None:
        if isinstance(ls, (int, float)) and ls > 100:
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = ls
        elif hasattr(ls, 'pt') and not isinstance(ls, (int, float)):
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = ls
        else:
            pf.line_spacing = ls

    if style_cfg.get("hanging_indent", 0) > 0:
        indent_pt = Pt(style_cfg["size"].pt * style_cfg["hanging_indent"])
        pf.left_indent = indent_pt
        pf.first_line_indent = Pt(-indent_pt.pt)
    elif style_cfg.get("first_line_indent", 0) > 0:
        pf.first_line_indent = Pt(style_cfg["size"].pt * style_cfg["first_line_indent"])
        pf.left_indent = None
    else:
        pf.first_line_indent = None
        pf.left_indent = None

    for run in para.runs:
        set_font(
            run,
            style_cfg["font_cn"],
            style_cfg["font_en"],
            style_cfg["size"],
            style_cfg.get("bold", False),
        )


def apply_three_line_table(table, table_cfg):
    """Apply three-line (三线表) formatting to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                f"</w:tcBorders>"
            )

            if row_idx == 0:
                tcBorders.find(qn("w:top")).set(qn("w:val"), "single")
                tcBorders.find(qn("w:top")).set(qn("w:sz"), "12")
                tcBorders.find(qn("w:bottom")).set(qn("w:val"), "single")
                tcBorders.find(qn("w:bottom")).set(qn("w:sz"), "6")
            elif row_idx == len(table.rows) - 1:
                tcBorders.find(qn("w:bottom")).set(qn("w:val"), "single")
                tcBorders.find(qn("w:bottom")).set(qn("w:sz"), "12")

            tcPr.append(tcBorders)

            for para in cell.paragraphs:
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    set_font(
                        run,
                        table_cfg["font_cn"],
                        table_cfg["font_en"],
                        table_cfg["size"],
                    )


def add_page_numbers(doc, size):
    """Add page numbers to footer (centered)."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        if not footer.paragraphs:
            para = footer.add_paragraph()
        else:
            para = footer.paragraphs[0]

        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._element.append(fldChar1)

        run2 = para.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._element.append(instrText)

        run3 = para.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._element.append(fldChar2)

        for r in [run, run2, run3]:
            r.font.size = size


_ROLE_TO_STYLE = {
    "title": "Title",
    "heading1": "Heading 1",
    "heading2": "Heading 2",
    "heading3": "Heading 3",
}


def _apply_word_style(para, role, doc):
    """Set the Word built-in style so that TOC and navigation pane work correctly."""
    style_name = _ROLE_TO_STYLE.get(role)
    if not style_name:
        return
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        pass


def format_document(input_path, output_path, template_key="通用论文", custom_template=None):
    """Main entry: format a Word document according to the chosen template."""
    doc = Document(input_path)
    tpl = custom_template if custom_template else TEMPLATES[template_key]

    total = len(doc.paragraphs)
    content_start = find_content_start(doc.paragraphs)

    content_section_idx = _find_content_section(doc, content_start)
    for idx, section in enumerate(doc.sections):
        if idx >= content_section_idx:
            section.top_margin = tpl["margins"]["top"]
            section.bottom_margin = tpl["margins"]["bottom"]
            section.left_margin = tpl["margins"]["left"]
            section.right_margin = tpl["margins"]["right"]

    changes = {
        "title": 0, "heading1": 0, "heading2": 0,
        "heading3": 0, "body": 0, "caption": 0,
        "reference": 0, "empty": 0, "skipped": 0,
    }

    in_reference_section = False

    for i, para in enumerate(doc.paragraphs):
        if i < content_start:
            changes["skipped"] += 1
            continue

        role = detect_role(para, i - content_start, total - content_start)

        if role == "heading1":
            text = para.text.strip()
            in_reference_section = bool(REFERENCE_HEADING_RE.match(text))
        elif role in ("heading2", "heading3"):
            pass
        elif in_reference_section and role == "body":
            role = "reference"

        changes[role] += 1

        if role == "empty":
            continue

        _apply_word_style(para, role, doc)

        if role in tpl:
            apply_paragraph_format(para, tpl[role])

    for table in doc.tables:
        if _table_is_in_preamble(table, doc, content_start):
            continue
        if tpl["table"].get("three_line"):
            apply_three_line_table(table, tpl["table"])
        else:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            set_font(
                                run,
                                tpl["table"]["font_cn"],
                                tpl["table"]["font_en"],
                                tpl["table"]["size"],
                            )

    if tpl.get("page_number"):
        add_page_numbers(doc, tpl["header_footer_size"])

    doc.save(output_path)

    report = _build_report(tpl, changes)

    return {
        "template": tpl["name"],
        "changes": changes,
        "report": report,
        "total_paragraphs": total,
        "content_start": content_start,
        "total_tables": len(doc.tables),
    }


def _find_content_section(doc, content_start):
    """Return the section index that the content-start paragraph belongs to."""
    if content_start == 0:
        return 0
    para_count = 0
    body_elements = doc.element.body
    section_idx = 0
    for child in body_elements:
        if child.tag.endswith("}p"):
            if para_count == content_start:
                return section_idx
            para_count += 1
        if child.tag.endswith("}sectPr"):
            section_idx += 1
    for p_elem in body_elements.iterchildren():
        if p_elem.tag.endswith("}p"):
            pPr = p_elem.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                section_idx += 1
    return 0


def _table_is_in_preamble(table, doc, content_start):
    """Check whether a table sits before the content-start paragraph in the document body."""
    if content_start == 0:
        return False
    body = doc.element.body
    table_elem = table._element
    para_count = 0
    for child in body:
        if child is table_elem:
            return para_count < content_start
        if child.tag.endswith("}p"):
            para_count += 1
    return False


_PT_TO_NAME = {v.pt: k for k, v in FONT_SIZE_MAP.items()}

_ALIGN_TO_CN = {
    WD_ALIGN_PARAGRAPH.CENTER: "居中",
    WD_ALIGN_PARAGRAPH.LEFT: "左对齐",
    WD_ALIGN_PARAGRAPH.RIGHT: "右对齐",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐",
}

_ROLE_CN = {
    "title": "题目",
    "heading1": "一级标题",
    "heading2": "二级标题",
    "heading3": "三级标题",
    "body": "正文",
    "caption": "图表标题",
    "reference": "参考文献",
}


def _format_line_spacing(ls):
    if ls is None:
        return "-"
    if isinstance(ls, (int, float)) and ls > 100:
        return f"固定值{ls / 12700:.0f}磅"
    if hasattr(ls, "pt") and not isinstance(ls, (int, float)):
        return f"固定值{ls.pt}磅"
    return f"{ls}倍"


def _build_report(tpl, changes):
    """Build a human-readable report dict summarising the format applied to each role."""
    report = {}
    for role in ("title", "heading1", "heading2", "heading3", "body", "caption", "reference"):
        count = changes.get(role, 0)
        if count == 0:
            continue
        cfg = tpl.get(role)
        if not cfg:
            continue
        size_pt = cfg["size"].pt if hasattr(cfg["size"], "pt") else cfg["size"]
        if isinstance(size_pt, (int, float)) and size_pt > 100:
            size_pt = size_pt / 12700

        indent_desc = ""
        if cfg.get("hanging_indent", 0) > 0:
            indent_desc = f"悬挂{cfg['hanging_indent']}字符"
        elif cfg.get("first_line_indent", 0) > 0:
            indent_desc = f"首行{cfg['first_line_indent']}字符"
        else:
            indent_desc = "无"

        report[role] = {
            "label": _ROLE_CN.get(role, role),
            "count": count,
            "font": cfg.get("font_cn", ""),
            "size": _PT_TO_NAME.get(size_pt, f"{size_pt}pt"),
            "bold": cfg.get("bold", False),
            "alignment": _ALIGN_TO_CN.get(cfg.get("alignment"), "左对齐"),
            "line_spacing": _format_line_spacing(cfg.get("line_spacing")),
            "indent": indent_desc,
        }
    return report
